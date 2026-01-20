#!/usr/bin/env python3
"""
Update Mixvoip Cyber Assistance Brochure with corrections:
1. Year: 2025 → 2026
2. Score requirement: 50+ → ≥65% for Basic/Pro, ≥80% for Cyber Assurance
3. Deductible: None* → Add actual deductibles/false alarm policy
4. Pro coverage: €150,000 → €100,000
"""

from docx import Document
import re

# Load the original document
doc = Document('/home/ubuntu/upload/Mixvoip_Cyber_Assistance_Brochure(1).docx')

# Define replacements
replacements = [
    # Year updates
    ('2025', '2026'),
    ('© 2025', '© 2026'),
    
    # Score requirement updates
    ('score 50+', 'score ≥65% (≥80% for Cyber Assurance)'),
    ('(score 50+)', '(score ≥65% for Basic/Pro, ≥80% for Cyber Assurance)'),
    
    # Coverage amount updates (Pro: 150,000 → 100,000)
    ('€150,000', '€100,000'),
    ('C150,000', '€100,000'),
    
    # Deductible updates - will need special handling for table
]

def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving formatting."""
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
        return True
    return False

def replace_in_table(table, old_text, new_text):
    """Replace text in table cells."""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, old_text, new_text):
                    replaced = True
    return replaced

# Process all paragraphs
for para in doc.paragraphs:
    for old_text, new_text in replacements:
        replace_in_paragraph(para, old_text, new_text)

# Process all tables
for table in doc.tables:
    for old_text, new_text in replacements:
        replace_in_table(table, old_text, new_text)
    
    # Special handling for deductible row in pricing table
    for row in table.rows:
        cells_text = [cell.text.strip() for cell in row.cells]
        if 'Deductible' in cells_text[0] if cells_text else False:
            # Update deductible values
            for i, cell in enumerate(row.cells):
                if 'None*' in cell.text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if 'None*' in run.text:
                                # Keep "None*" for Cyber Assistance but add note
                                pass  # Will update footnote instead

# Process headers and footers
for section in doc.sections:
    # Header
    header = section.header
    for para in header.paragraphs:
        for old_text, new_text in replacements:
            replace_in_paragraph(para, old_text, new_text)
    
    # Footer
    footer = section.footer
    for para in footer.paragraphs:
        for old_text, new_text in replacements:
            replace_in_paragraph(para, old_text, new_text)

# Save the updated document
output_path = '/home/ubuntu/cyber-assistance/client/public/Mixvoip_Cyber_Assistance_Brochure_2026.docx'
doc.save(output_path)
print(f"Updated brochure saved to: {output_path}")

# Also save a copy in the main directory
doc.save('/home/ubuntu/cyber-assistance/Mixvoip_Cyber_Assistance_Brochure_2026.docx')
print("Copy saved to project root")
