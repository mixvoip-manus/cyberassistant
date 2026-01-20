#!/usr/bin/env python3
"""
Update Mixvoip Cyber Assistance Brochure - Version 2
Fix remaining issues:
1. Score requirement: 50+ → ≥65% for Basic/Pro, ≥80% for Cyber Assurance
2. Update deductible footnote
"""

from docx import Document
from docx.shared import Pt
import re

# Load the document
doc = Document('/home/ubuntu/cyber-assistance/Mixvoip_Cyber_Assistance_Brochure_2026.docx')

def replace_text_in_runs(paragraph, old_text, new_text):
    """Replace text across runs in a paragraph."""
    full_text = paragraph.text
    if old_text in full_text:
        # Find and replace in the full paragraph text
        new_full_text = full_text.replace(old_text, new_text)
        
        # Clear all runs and add new text to first run
        if paragraph.runs:
            # Store formatting from first run
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold
            
            # Clear all runs
            for run in paragraph.runs:
                run.text = ""
            
            # Set new text in first run
            paragraph.runs[0].text = new_full_text
        return True
    return False

def process_document(doc):
    """Process all paragraphs in the document."""
    replacements = [
        # Score requirement - various formats
        ('score 50+', 'score ≥65% (≥80% for Cyber Assurance)'),
        ('(score 50+)', '(score ≥65% for Basic/Pro, ≥80% for Cyber Assurance)'),
        
        # Deductible footnote
        ('* No deductible when subscribing online.', 
         '* Cyber Assistance: 2 false alarms/year free (Basic), then 50€/false alarm. Cyber Assurance deductible: 500€ (50k), 1,000€ (100k), 2,500€ (250k).'),
        ('No deductible when subscribing online', 
         'Cyber Assistance: 2 false alarms/year free (Basic), then 50€/false alarm. Cyber Assurance deductible: 500€ (50k), 1,000€ (100k), 2,500€ (250k)'),
    ]
    
    # Process all paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements:
            if old_text in para.text:
                replace_text_in_runs(para, old_text, new_text)
    
    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements:
                        if old_text in para.text:
                            replace_text_in_runs(para, old_text, new_text)
    
    # Process headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            for old_text, new_text in replacements:
                if old_text in para.text:
                    replace_text_in_runs(para, old_text, new_text)
        
        for para in section.footer.paragraphs:
            for old_text, new_text in replacements:
                if old_text in para.text:
                    replace_text_in_runs(para, old_text, new_text)

# Process the document
process_document(doc)

# Save
output_path = '/home/ubuntu/cyber-assistance/client/public/Mixvoip_Cyber_Assistance_Brochure_2026.docx'
doc.save(output_path)
print(f"Updated brochure saved to: {output_path}")

# Also save a copy
doc.save('/home/ubuntu/cyber-assistance/Mixvoip_Cyber_Assistance_Brochure_2026.docx')
print("Copy saved to project root")
